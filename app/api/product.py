from pydantic import BaseModel
from typing import List, Dict, Any
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.encoders import jsonable_encoder
from io import BytesIO
from fastapi.responses import StreamingResponse
from docx import Document

from app.services.usage_service import (
    get_usage_snapshot,
    increment_usage,
    enforce_usage_limit,
)
from app.services.citation_service import format_all_citations
from app.services.workspace_service import (
    add_favorite as add_favorite_local,
    add_draft_section,
    add_note as add_note_local,
    add_saved_paper as add_saved_paper_local,
    add_to_reading_queue as add_to_reading_queue_local,
    get_workspace as get_local_workspace,
    remove_saved_paper as remove_saved_paper_local,
)
from app.services.persistence_service import (
    create_or_update_profile,
    get_profile_by_clerk_user_id,
    get_saved_papers,
    save_paper,
    delete_saved_paper,
    get_reading_queue,
    queue_paper,
    get_favorites,
    favorite_paper,
    add_note as persist_note,
    get_comparisons,
    save_comparison as persist_comparison,
    delete_comparison as persist_delete_comparison,
)
from app.services.authoring_service import (
    create_project,
    list_projects,
    get_project,
    add_section,
    add_source_to_project,
    replace_working_draft,
    update_project_abstract,
    render_project_markdown,
    render_project_html,
)
from app.services.upload_service import (
    ingest_local_file,
    ingest_url,
)
from app.services.recommendation_tracking_service import (
    safe_track_recommendation_activity,
    safe_track_recommendation_activity_for_user,
)
from app.services.business_ops_service import sync_profile_to_business_records
from app.services.stripe_checkout_service import StripeConfigError, create_checkout_session
from app.services.clerk_auth_service import require_verified_clerk_user
from app.services.local_state_service import (
    create_or_update_local_profile,
    delete_local_comparison,
    get_local_comparisons,
    get_local_profile,
    save_local_comparison,
)

class ComparisonHistoryItem(BaseModel):
    title: str
    work_ids: List[str]
    paper_titles: List[str]
    question: str | None = None
    summary: str | None = None


class ClerkProfileSyncRequest(BaseModel):
    clerk_user_id: str
    username: str | None = None
    email: str | None = None
    full_name: str | None = None
    first_name: str | None = None
    last_name: str | None = None
    avatar_url: str | None = None
    institution: str | None = None
    role_title: str | None = None
    github_url: str | None = None
    linkedin_url: str | None = None
    research_interests: List[str] | None = None
    onboarding_notes: str | None = None
    plan: str = "free"
    auth_provider: str = "clerk"


class CheckoutSessionRequest(BaseModel):
    plan_code: str
    clerk_user_id: str
    email: str | None = None
    success_url: str | None = None
    cancel_url: str | None = None


_UPLOADED_SOURCES: Dict[str, List[Dict[str, Any]]] = {}
_UPLOADED_DOCUMENTS: Dict[str, Dict[str, Any]] = {}


router = APIRouter()


def resolve_profile_id(user_id: str) -> str:
    profile = get_profile_by_clerk_user_id(user_id)
    if not profile:
        raise HTTPException(status_code=404, detail="Profile not found")
    return profile["id"]


def require_current_user(requested_user_id: str, identity: dict) -> str:
    clerk_user_id = identity.get("clerk_user_id")
    if requested_user_id != clerk_user_id:
        raise HTTPException(status_code=403, detail="You do not have access to this user resource.")
    return clerk_user_id


def get_effective_plan(clerk_user_id: str) -> str:
    profile = None
    try:
        profile = get_profile_by_clerk_user_id(clerk_user_id)
    except Exception:
        profile = None
    if not profile:
        profile = get_local_profile(clerk_user_id)
    return (profile or {}).get("plan") or "free"


def _local_workspace_item(items: list[dict], work_id: str | None) -> dict:
    if not items:
        return {}
    if work_id is None:
        return items[-1]
    for item in items:
        if item.get("work_id") == work_id:
            return item
    return items[-1]


def _register_uploaded_document(file_name: str, result: dict, source_type: str, owner_user_id: str) -> str:
    document_id = f"uploaded_{len(_UPLOADED_DOCUMENTS) + 1}"

    text_value = (
        result.get("text")
        or result.get("content")
        or result.get("raw_text")
        or result.get("document_text")
        or ""
    )

    summary_value = (
        result.get("summary")
        or result.get("plain_english_summary")
        or result.get("executive_summary")
        or ""
    )

    metadata = {
        "document_id": document_id,
        "work_id": document_id,
        "title": file_name or "Uploaded Document",
        "author": "User Upload",
        "institution": "Local Upload",
        "topic": "User Provided",
        "primary_topic": "User Provided",
        "source_system": "uploaded",
        "publication_year": datetime.utcnow().year,
        "published": datetime.utcnow().date().isoformat(),
        "citation": f"User Upload ({datetime.utcnow().year}). {file_name or 'Uploaded Document'}.",
        "availability_label": "Full Document",
        "has_full_document": True,
        "pdf_url": None,
        "entry_url": None,
        "categories": source_type,
    }

    ai_summary = {
        "plain_english_summary": summary_value or "No summary available.",
        "academic_summary": summary_value or "No academic summary available.",
        "methods_summary": result.get("methods_summary") or "Not extracted yet.",
        "results_summary": result.get("results_summary") or "Not extracted yet.",
        "limitations_summary": result.get("limitations_summary") or "Not extracted yet.",
        "practical_applications": result.get("practical_applications") or "Not extracted yet.",
        "suggested_topics": result.get("suggested_topics") or "",
        "citation_guidance": result.get("citation_guidance") or "",
    }

    _UPLOADED_DOCUMENTS[document_id] = {
        "metadata": metadata,
        "ai_summary": ai_summary,
        "content": text_value,
        "source_type": source_type,
        "owner_user_id": owner_user_id,
        "created_at": datetime.utcnow().isoformat(),
    }

    return document_id


@router.get("/health")
def product_health():
    return {"status": "product api ok"}


# -------------------------
# Usage
# -------------------------
@router.get("/usage/{user_id}")
def usage_snapshot(user_id: str, identity: dict = Depends(require_verified_clerk_user)):
    clerk_user_id = require_current_user(user_id, identity)
    return jsonable_encoder(get_usage_snapshot(clerk_user_id, get_effective_plan(clerk_user_id)))


# -------------------------
# Citations
# -------------------------
@router.post("/citations")
def citations(payload: dict):
    return jsonable_encoder(format_all_citations(payload))


@router.post("/auth/sync-profile")
def sync_profile(payload: ClerkProfileSyncRequest, identity: dict = Depends(require_verified_clerk_user)):
    clerk_user_id = identity.get("clerk_user_id")
    if payload.clerk_user_id != clerk_user_id:
        raise HTTPException(status_code=403, detail="Profile sync identity mismatch.")
    try:
        existing = get_profile_by_clerk_user_id(clerk_user_id) or get_local_profile(clerk_user_id) or {}
    except Exception:
        existing = get_local_profile(clerk_user_id) or {}
    safe_payload = payload.model_dump(exclude_none=True)
    safe_payload["clerk_user_id"] = clerk_user_id
    safe_payload["plan"] = existing.get("plan") or "free"
    try:
        profile = create_or_update_profile(safe_payload)
        try:
            business_records = sync_profile_to_business_records(
                profile,
                acquisition_context={
                    "conversion_source": "product.auth_sync",
                    "metadata": {
                        "requested_plan": payload.plan,
                        "effective_plan": profile.get("plan", "free"),
                        "auth_provider": payload.auth_provider,
                    },
                },
            )
        except Exception:
            business_records = {"customer": None, "sales_profile": None, "mode": "local_fallback"}
        return jsonable_encoder({
            "success": True,
            "profile": profile,
            **business_records,
        })
    except Exception as exc:
        profile = create_or_update_local_profile(safe_payload)
        return jsonable_encoder({
            "success": True,
            "profile": profile,
            "customer": {
                "profile_id": profile["id"],
                "clerk_user_id": profile["clerk_user_id"],
                "plan_code": profile.get("plan", "free"),
            },
            "sales_profile": {
                "current_plan": profile.get("plan", "free"),
                "recommended_upgrade_plan": None,
                "top_topics": [],
            },
            "mode": "local_fallback",
            "detail": f"Supabase sync unavailable: {exc}",
        })


@router.post("/billing/checkout-session")
def billing_checkout_session(payload: CheckoutSessionRequest, identity: dict = Depends(require_verified_clerk_user)):
    clerk_user_id = identity.get("clerk_user_id")
    if payload.clerk_user_id != clerk_user_id:
        raise HTTPException(status_code=403, detail="Checkout identity mismatch.")
    try:
        profile = get_profile_by_clerk_user_id(clerk_user_id) or get_local_profile(clerk_user_id) or {}
    except Exception:
        profile = get_local_profile(clerk_user_id) or {}
    try:
        session = create_checkout_session(
            plan_code=payload.plan_code,
            clerk_user_id=clerk_user_id,
            email=profile.get("email") or payload.email,
            success_url=payload.success_url,
            cancel_url=payload.cancel_url,
        )
        return jsonable_encoder({
            "success": True,
            "checkout_url": session.get("url"),
            "checkout_session_id": session.get("id"),
            "mode": "stripe",
        })
    except StripeConfigError as exc:
        return jsonable_encoder({
            "success": False,
            "mode": "unconfigured",
            "detail": str(exc),
        })
    except Exception as exc:
        return jsonable_encoder({
            "success": False,
            "mode": "error",
            "detail": f"Stripe checkout creation failed: {exc}",
        })


# -------------------------
# Workspace
# -------------------------
@router.get("/workspace/{user_id}")
def workspace(user_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        return jsonable_encoder({
            "user_id": user_id,
            "saved_papers": get_saved_papers(profile_id),
            "reading_queue": get_reading_queue(profile_id),
            "favorites": get_favorites(profile_id),
            "mode": "supabase",
        })
    except Exception:
        local_profile = get_local_profile(user_id) or create_or_update_local_profile({"clerk_user_id": user_id, "plan": "free"})
        local_workspace = get_local_workspace(user_id)
        return jsonable_encoder({
            "user_id": user_id,
            "saved_papers": local_workspace.get("saved_papers", []),
            "reading_queue": local_workspace.get("reading_queue", []),
            "favorites": local_workspace.get("favorites", []),
            "notes": local_workspace.get("notes", []),
            "profile": local_profile,
            "mode": "local_fallback",
        })


@router.post("/workspace/{user_id}/save-paper")
def workspace_save_paper(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        result = save_paper(profile_id, payload)
    except Exception:
        create_or_update_local_profile({"clerk_user_id": user_id, "plan": get_effective_plan(user_id)})
        workspace = add_saved_paper_local(user_id, payload)
        result = _local_workspace_item(workspace.get("saved_papers", []), payload.get("work_id"))
        profile_id = f"local_{user_id}"
    safe_track_recommendation_activity(
        profile_id,
        "save",
        {
            **payload,
            "event_source": "product.workspace",
            "action_context": "save_paper",
        },
    )
    return jsonable_encoder(result)


@router.post("/workspace/{user_id}/queue-paper")
def workspace_queue_paper(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        result = queue_paper(profile_id, payload)
    except Exception:
        create_or_update_local_profile({"clerk_user_id": user_id, "plan": get_effective_plan(user_id)})
        workspace = add_to_reading_queue_local(user_id, payload)
        result = _local_workspace_item(workspace.get("reading_queue", []), payload.get("work_id"))
        profile_id = f"local_{user_id}"
    safe_track_recommendation_activity(
        profile_id,
        "queue",
        {
            **payload,
            "event_source": "product.workspace",
            "action_context": "queue_paper",
        },
    )
    return jsonable_encoder(result)


@router.post("/workspace/{user_id}/favorite-paper")
def workspace_favorite_paper(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        result = favorite_paper(profile_id, payload)
    except Exception:
        create_or_update_local_profile({"clerk_user_id": user_id, "plan": get_effective_plan(user_id)})
        workspace = add_favorite_local(user_id, payload)
        result = _local_workspace_item(workspace.get("favorites", []), payload.get("work_id"))
        profile_id = f"local_{user_id}"
    safe_track_recommendation_activity(
        profile_id,
        "favorite",
        {
            **payload,
            "event_source": "product.workspace",
            "action_context": "favorite_paper",
        },
    )
    return jsonable_encoder(result)


@router.post("/workspace/{user_id}/note")
def workspace_note(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        result = persist_note(profile_id, payload)
    except Exception:
        create_or_update_local_profile({"clerk_user_id": user_id, "plan": "free"})
        workspace = add_note_local(user_id, payload)
        result = workspace.get("notes", [])[-1]
        profile_id = f"local_{user_id}"
    safe_track_recommendation_activity(
        profile_id,
        "note",
        {
            **payload,
            "work_id": payload.get("paper_work_id"),
            "document_id": payload.get("paper_document_id"),
            "title": payload.get("paper_title"),
            "event_source": "product.workspace",
            "action_context": "paper_note",
        },
    )
    return jsonable_encoder(result)


@router.post("/workspace/{user_id}/draft-section")
def workspace_draft_section(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    result = add_draft_section(user_id, payload)
    return jsonable_encoder(result)


@router.delete("/workspace/{user_id}/paper/{work_id}")
def workspace_remove_paper(user_id: str, work_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        delete_saved_paper(profile_id, work_id)
    except Exception:
        remove_saved_paper_local(user_id, work_id)
    return jsonable_encoder({
        "success": True,
        "work_id": work_id,
    })

@router.get("/workspace/{user_id}/comparisons")
def get_saved_comparisons(user_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        comparisons = get_comparisons(profile_id)
        mode = "supabase"
    except Exception:
        comparisons = get_local_comparisons(user_id)
        mode = "local_fallback"
    return jsonable_encoder({
        "user_id": user_id,
        "comparisons": comparisons,
        "mode": mode,
    })


@router.post("/workspace/{user_id}/comparisons")
def save_comparison(user_id: str, payload: ComparisonHistoryItem, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        item = persist_comparison(profile_id, payload.model_dump())
    except Exception:
        item = save_local_comparison(user_id, payload.model_dump())
        profile_id = f"local_{user_id}"
    for index, work_id in enumerate(payload.work_ids):
        safe_track_recommendation_activity(
            profile_id,
            "compare",
            {
                "work_id": work_id,
                "title": payload.paper_titles[index] if index < len(payload.paper_titles) else None,
                "event_source": "product.workspace",
                "action_context": "comparison_saved",
                "metadata": {
                    "comparison_id": item.get("id"),
                    "comparison_title": payload.title,
                    "question": payload.question,
                    "all_work_ids": payload.work_ids,
                },
            },
        )
    return jsonable_encoder({
        "success": True,
        "comparison": item
    })


@router.delete("/workspace/{user_id}/comparisons/{comparison_id}")
def delete_comparison(user_id: str, comparison_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    try:
        profile_id = resolve_profile_id(user_id)
        persist_delete_comparison(profile_id, comparison_id)
    except Exception:
        delete_local_comparison(user_id, comparison_id)
    return jsonable_encoder({
        "success": True,
        "comparison_id": comparison_id
    })


# -------------------------
# Authoring
# -------------------------
@router.post("/authoring/{user_id}/projects")
def authoring_create_project(user_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    title = payload.get("title", "Untitled Project")
    project_type = payload.get("project_type", "white_paper")
    result = create_project(user_id, title, project_type)
    return jsonable_encoder(result)


@router.get("/authoring/{user_id}/projects")
def authoring_list_projects(user_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    return jsonable_encoder(list_projects(user_id))


@router.get("/authoring/{user_id}/projects/{project_id}")
def authoring_get_project(user_id: str, project_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    project = get_project(user_id, project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return jsonable_encoder(project)


@router.post("/authoring/{user_id}/projects/{project_id}/section")
def authoring_add_section(user_id: str, project_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    title = payload.get("title", "Untitled Section")
    content = payload.get("content", "")
    result = add_section(user_id, project_id, title, content)
    return jsonable_encoder(result)


@router.post("/authoring/{user_id}/projects/{project_id}/source")
def authoring_add_source(user_id: str, project_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    result = add_source_to_project(user_id, project_id, payload)
    return jsonable_encoder(result)


@router.post("/authoring/{user_id}/projects/{project_id}/abstract")
def authoring_set_abstract(user_id: str, project_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    abstract = payload.get("abstract", "")
    result = update_project_abstract(user_id, project_id, abstract)
    return jsonable_encoder(result)


@router.post("/authoring/{user_id}/projects/{project_id}/draft")
def authoring_replace_draft(user_id: str, project_id: str, payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    result = replace_working_draft(
        user_id,
        project_id,
        payload.get("title", "Untitled Project"),
        payload.get("abstract", ""),
        payload.get("content", ""),
    )
    return jsonable_encoder(result)


@router.get("/authoring/{user_id}/projects/{project_id}/render/markdown")
def authoring_render_markdown(user_id: str, project_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    return {"markdown": render_project_markdown(user_id, project_id)}


@router.get("/authoring/{user_id}/projects/{project_id}/render/html")
def authoring_render_html(user_id: str, project_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    return {"html": render_project_html(user_id, project_id)}

@router.get("/authoring/{user_id}/projects/{project_id}/export/docx")
def export_project_docx(user_id: str, project_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    project = get_project(user_id, project_id)

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    doc = Document()

    title = project.get("title") or "Untitled Project"
    abstract = project.get("abstract") or ""
    sections = project.get("sections", []) or []
    sources = project.get("sources", []) or []

    doc.add_heading(title, level=0)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract or "No abstract provided.")

    if sections:
        doc.add_heading("Sections", level=1)
        for section in sections:
            section_title = section.get("title") or "Untitled Section"
            section_content = section.get("content") or ""
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(section_content or "No content provided.")
    else:
        doc.add_heading("Sections", level=1)
        doc.add_paragraph("No sections added.")

    doc.add_heading("References", level=1)
    if sources:
        for source in sources:
            citation = source.get("citation") or source.get("title") or "Untitled Source"
            doc.add_paragraph(citation, style="List Bullet")
    else:
        doc.add_paragraph("No references added.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{title.strip().replace(' ', '_') or 'research_project'}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )

@router.get("/authoring/{user_id}/projects/{project_id}/export/pdf")
def export_project_pdf(user_id: str, project_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    project = get_project(user_id, project_id)

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.utils import simpleSplit
        from reportlab.pdfgen import canvas
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="PDF export requires the reportlab package to be installed."
        ) from exc

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER

    left_margin = 50
    right_margin = 50
    top_margin = 50
    bottom_margin = 50
    usable_width = width - left_margin - right_margin
    y = height - top_margin

    title = project.get("title") or "Untitled Project"
    abstract = project.get("abstract") or ""
    sections = project.get("sections", []) or []
    sources = project.get("sources", []) or []

    def new_page_if_needed(current_y: float, needed: float = 24) -> float:
        nonlocal pdf
        if current_y - needed < bottom_margin:
            pdf.showPage()
            return height - top_margin
        return current_y

    def write_wrapped(text: str, font_name: str = "Helvetica", font_size: int = 11, gap: int = 16) -> None:
        nonlocal y
        text = text or ""
        lines = simpleSplit(text, font_name, font_size, usable_width)
        pdf.setFont(font_name, font_size)
        for line in lines:
            y = new_page_if_needed(y, gap)
            pdf.drawString(left_margin, y, line)
            y -= gap

    pdf.setTitle(title)

    y = new_page_if_needed(y, 30)
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(left_margin, y, title)
    y -= 30

    y = new_page_if_needed(y, 22)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left_margin, y, "Abstract")
    y -= 20
    write_wrapped(abstract or "No abstract provided.")
    y -= 8

    y = new_page_if_needed(y, 22)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left_margin, y, "Sections")
    y -= 20

    if sections:
        for section in sections:
            section_title = section.get("title") or "Untitled Section"
            section_content = section.get("content") or "No content provided."

            y = new_page_if_needed(y, 22)
            pdf.setFont("Helvetica-Bold", 12)
            pdf.drawString(left_margin, y, section_title)
            y -= 18
            write_wrapped(section_content)
            y -= 8
    else:
        write_wrapped("No sections added.")
        y -= 8

    y = new_page_if_needed(y, 22)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(left_margin, y, "References")
    y -= 20

    if sources:
        for source in sources:
            citation = source.get("citation") or source.get("title") or "Untitled Source"
            write_wrapped(f"• {citation}")
    else:
        write_wrapped("No references added.")

    pdf.save()
    buffer.seek(0)

    filename = f"{title.strip().replace(' ', '_') or 'research_project'}.pdf"

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        },
    )


# -------------------------
# Uploads
# -------------------------
@router.post("/uploads/local")
async def upload_local_file(
    user_id: str = Form(...),
    file: UploadFile = File(...),
    identity: dict = Depends(require_verified_clerk_user),
):
    user_id = require_current_user(user_id, identity)
    plan = get_effective_plan(user_id)
    try:
        enforce_usage_limit(user_id, "uploads", plan)
    except PermissionError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc

    import os
    import tempfile

    suffix = os.path.splitext(file.filename or "")[1]
    tmp_path = None

    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="A file is required")

        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            content = await file.read()
            if not content:
                raise HTTPException(status_code=400, detail="The selected file is empty")
            tmp.write(content)
            tmp_path = tmp.name

        result = ingest_local_file(tmp_path)
        document_id = _register_uploaded_document(file.filename or "uploaded_file", result, "local_file", user_id)
        increment_usage(user_id, "uploads", 1)

        user_uploads = _UPLOADED_SOURCES.setdefault(user_id, [])

        item = {
            "upload_id": f"local_{len(user_uploads) + 1}",
            "source_type": "local_file",
            "file_name": result.get("file_name") or file.filename or "uploaded_file",
            "document_id": document_id,
            "file_path": result.get("file_path"),
            "status": "ingested",
            "created_at": datetime.utcnow().isoformat(),
            "result": result,
        }

        user_uploads.insert(0, item)

        safe_track_recommendation_activity_for_user(
            user_id,
            "upload",
            {
                **(_UPLOADED_DOCUMENTS.get(document_id, {}).get("metadata") or {}),
                "document_id": document_id,
                "title": item["file_name"],
                "event_source": "product.uploads",
                "action_context": "local_upload",
                "metadata": {
                    "upload_id": item["upload_id"],
                    "source_type": item["source_type"],
                },
            },
        )

        return jsonable_encoder({
            "success": True,
            "upload": item,
            "file_name": item["file_name"],
            "result": result,
        })
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Local upload failed: {exc}") from exc

    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


@router.post("/uploads/url")
def upload_url(payload: dict, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(payload.get("user_id", ""), identity)
    plan = get_effective_plan(user_id)
    url = (payload.get("url", "") or "").strip()

    if not url:
        raise HTTPException(status_code=400, detail="URL is required")

    if not url.startswith(("http://", "https://")):
        url = f"https://{url}"

    try:
        enforce_usage_limit(user_id, "uploads", plan)
    except PermissionError as exc:
        raise HTTPException(status_code=403, detail=str(exc)) from exc

    try:
        result = ingest_url(url)
        document_id = _register_uploaded_document(result.get("file_name") or url.split("/")[-1] or "web_document", result, "url", user_id)
        increment_usage(user_id, "uploads", 1)

        user_uploads = _UPLOADED_SOURCES.setdefault(user_id, [])

        item = {
            "upload_id": f"url_{len(user_uploads) + 1}",
            "source_type": "url",
            "file_name": result.get("file_name") or url.split("/")[-1] or "web_document",
            "document_id": document_id,
            "url": url,
            "status": "ingested",
            "created_at": datetime.utcnow().isoformat(),
            "result": result,
        }

        user_uploads.insert(0, item)

        safe_track_recommendation_activity_for_user(
            user_id,
            "upload",
            {
                **(_UPLOADED_DOCUMENTS.get(document_id, {}).get("metadata") or {}),
                "document_id": document_id,
                "title": item["file_name"],
                "event_source": "product.uploads",
                "action_context": "url_upload",
                "metadata": {
                    "upload_id": item["upload_id"],
                    "source_type": item["source_type"],
                    "url": url,
                },
            },
        )

        return jsonable_encoder({
            "success": True,
            "upload": item,
            "file_name": item["file_name"],
            "result": result,
        })
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"URL ingest failed: {exc}") from exc


@router.get("/uploads/{user_id}")
def list_uploads(user_id: str, identity: dict = Depends(require_verified_clerk_user)):
    user_id = require_current_user(user_id, identity)
    return jsonable_encoder({
        "user_id": user_id,
        "uploads": _UPLOADED_SOURCES.get(user_id, [])
    })


@router.get("/uploads/document/{document_id}")
def get_uploaded_document(document_id: str, identity: dict = Depends(require_verified_clerk_user)):
    doc = _UPLOADED_DOCUMENTS.get(document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Uploaded document not found")
    if doc.get("owner_user_id") != identity.get("clerk_user_id"):
        raise HTTPException(status_code=403, detail="You do not have access to this uploaded document.")
    return jsonable_encoder(doc)
